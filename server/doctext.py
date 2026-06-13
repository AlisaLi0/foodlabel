"""从上传的标签文档（txt/markdown/csv/pdf/docx）提取纯文本.

当用户直接上传**已是文字**的标签内容（而非图片）时，无需 OCR——本模块把文档
解析成纯文本，交给后续识读/规则/合规判定流程，从而跳过识图步骤、更快更准。

支持：
  * .txt / .md / .csv 及任意 text/* —— 按常见中文编码解码
  * .pdf —— pypdf 提取文本层（扫描件/图片型 PDF 无文本层，提示改用图片走 OCR）
  * .docx —— python-docx 提取段落与表格（营养成分表常在表格里）

不支持 .doc（旧二进制 Word），提示用户另存为 .docx 或 PDF。解析失败统一抛 DocError，
上层（core.prepare_inputs）将其转为 InputError → HTTP 400。
"""
from __future__ import annotations

import io

# 文档扩展名（供前后端一致判断）。图片不在此列。
DOC_EXTS = (".txt", ".md", ".csv", ".pdf", ".docx")

_DOC_CONTENT_TYPES = {
    "text/plain",
    "text/markdown",
    "text/csv",
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",  # .doc（实际不支持，extract_text 里单独报错）
}


class DocError(ValueError):
    """文档解析失败或类型不支持。上层映射为 HTTP 400。"""


def is_doc(content_type: str | None, filename: str | None) -> bool:
    """按 content-type 或文件扩展名判断是否为受支持/可识别的文档（含 .doc）。"""
    name = (filename or "").lower()
    ct = (content_type or "").lower()
    if name.endswith(DOC_EXTS) or name.endswith(".doc"):
        return True
    if ct in _DOC_CONTENT_TYPES or ct.startswith("text/"):
        return True
    return False


def _decode_text(raw: bytes) -> str:
    """按常见中文编码尽力解码纯文本文件。"""
    for enc in ("utf-8-sig", "utf-8", "gb18030", "utf-16", "latin-1"):
        try:
            return raw.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return raw.decode("utf-8", errors="ignore")


def _extract_pdf(raw: bytes) -> str:
    """提取 PDF 文本层（逐页）。无文本层（扫描件）返回空串，由上层提示改用图片。"""
    try:
        from pypdf import PdfReader
    except ImportError as e:  # pragma: no cover - 依赖缺失时给明确提示
        raise DocError("服务器未安装 PDF 解析库（pypdf），无法处理 PDF。") from e
    try:
        reader = PdfReader(io.BytesIO(raw))
    except Exception as e:  # noqa: BLE001 - 损坏/加密 PDF
        raise DocError(f"PDF 无法解析：{e}") from e
    parts: list[str] = []
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception:  # noqa: BLE001 - 单页失败不致命
            t = ""
        if t.strip():
            parts.append(t.strip())
    return "\n".join(parts)


def _extract_docx(raw: bytes) -> str:
    """提取 .docx 的段落与表格文本（营养成分表通常在表格中）。"""
    try:
        import docx  # python-docx
    except ImportError as e:  # pragma: no cover
        raise DocError("服务器未安装 Word 解析库（python-docx），无法处理 .docx。") from e
    try:
        document = docx.Document(io.BytesIO(raw))
    except Exception as e:  # noqa: BLE001
        raise DocError(f".docx 无法解析：{e}") from e
    parts: list[str] = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    return "\n".join(parts)


def extract_text(raw: bytes, content_type: str | None, filename: str | None) -> str:
    """从文档字节提取纯文本。失败/不支持抛 DocError。"""
    name = (filename or "").lower()
    ct = (content_type or "").lower()

    if name.endswith(".doc") and not name.endswith(".docx"):
        raise DocError(".doc 旧格式暂不支持，请另存为 .docx 或 PDF 后再上传。")

    if name.endswith(".pdf") or ct == "application/pdf":
        text = _extract_pdf(raw)
        if not text.strip():
            raise DocError("该 PDF 没有可提取的文字（可能是扫描件/图片型 PDF），请改用图片上传走 OCR。")
        return text

    if name.endswith(".docx") or "wordprocessingml" in ct:
        text = _extract_docx(raw)
        if not text.strip():
            raise DocError("该 Word 文档未提取到文字内容。")
        return text

    if name.endswith((".txt", ".md", ".csv")) or ct.startswith("text/"):
        return _decode_text(raw)

    raise DocError(f"不支持的文件类型：{filename or content_type or '未知'}")
